#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed Mar 20 09:58:18 2019

@author: shridhar
"""

from docx import Document

document = Document('/home/shridhar/Downloads/PIN-BHARCS.docx')
    




import docx

doc = docx.Document('/home/shridhar/Downloads/PIN-BHARCS.docx')
print(len(doc.paragraphs))

print (doc.paragraphs[5].text)
doc.namelist()

import docx2txt
text = docx2txt.process('/home/shridhar/Downloads/PIN-BHARCS.docx')

print(text)

import re
result = re.findall(r'[^\w](\w{3})[^\w]', "Sajit Gopalan")
result
text[10]

print (text.index("Sajit Gopalan"))
print (text.index("Srinivasan Balaji"))
print (text.index("Deliverables"))
print (text.index("AWS"))
print (text.index("Abhishek L"))
print (text.index("Akhila U"))
print (text.index("PMO"))

letter=text[93]

import pandas as pd
text1=pd.DataFrame(text)
df = pd.DataFrame([text])
df[0]
df.rename(columns={0:'string'},inplace=True)




txt=["Sajit Gopalan","Srinivasan Balaji","Deliverables","AWS","Abhishek L","Akhila U","PMO"]
df['string'].str.contains('Sajit Gopalan')
import re
df1=re.findall("Sajit Gopalan", text)
df2=re.findall("Srinivasan Balaji", text)
df3=re.findall("Deliverables", text)
df4=re.findall("AWS", text)
df5=re.findall("Abhishek L", text)
df6=re.findall("Akhila U", text)
df7=re.findall("PMO", text)
list=[df1,df2,df3,df4,df5,df6,df7]
df9 = pd.DataFrame(list)
df9.to_csv('/home/shridhar/Desktop/picked_words_using_python.csv')

df1=re.findall("PIN ID  EBRC41S", text)
df2=re.findall("% ALLOCATION", text)
df10=re.findall("Abhishek L", text)
df9=re.findall("Akhila U", text)
df3=re.findall("START DATE", text)
df4=re.findall("END DATE", text)
df5=re.findall("04-12-2018", text)
df6=re.findall("31-0302109", text)
df7=re.findall("50%", text)
df11=re.findall("31-01-2019", text)
df8=re.findall("100%", text)

list=[df1,df2,df3,df4,df5,df6,df7,df8,df9,df10,df11]
df = pd.DataFrame(list)
df.to_csv('/home/shridhar/Desktop/picked_words_using_python.csv')
